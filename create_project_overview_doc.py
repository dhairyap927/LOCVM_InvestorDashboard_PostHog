from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUTPUT = Path("docs/LOCVM_PostHog_Project_Overview.docx")


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "light_fill": "F2F4F7",
    "mid_fill": "E8EEF5",
    "border": "D9E2EC",
    "muted": "5B677A",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=COLORS["border"]):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_column_widths(table, widths):
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill=COLORS["mid_fill"]):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
        if row_index == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(COLORS["ink"])


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    set_column_widths(table, widths)
    style_table(table)
    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FAFC")
    set_cell_border(cell, "B8C7D9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    title_paragraph = cell.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    body_paragraph = cell.add_paragraph()
    body_paragraph.add_run(body)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(document):
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("LOCVM PostHog Analytics Overview")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def build_document():
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("LOCVM PostHog Analytics Project Overview")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(COLORS["ink"])

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Standalone analytics, export, and investor reporting layer for LOCVM marketplace activity")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_callout(
        document,
        "Purpose",
        "This project turns LOCVM marketplace activity into clean PostHog analytics data, internal analyst exports, and investor-ready Excel reports. It is isolated from the existing GA4/Shiny dashboard."
    )

    document.add_heading("1. What the Project Does", level=1)
    add_bullets(document, [
        "Defines LOCVM business events for PostHog.",
        "Maps LOCVM user and locum fields into clean analytics properties.",
        "Exports raw PostHog event data to CSV.",
        "Uses an Excel-editable metric definition file to generate a metric matrix.",
        "Generates an investor-ready Excel workbook from event data.",
        "Supports mock testing before connecting real PostHog credentials.",
    ])

    document.add_heading("2. Project Workflow", level=1)
    workflow_rows = [
        ("1", "LOCVM app event happens", "A user signs up, searches locums, applies, reserves, completes a locum, or creates a clinic posting."),
        ("2", "Event is tracked", "The central analytics utility sends a clean PostHog event with normalized fields."),
        ("3", "Events are exported", "The export script retrieves PostHog events for a selected time period."),
        ("4", "Metrics are calculated", "The metric matrix script calculates counts, unique users, sums, averages, filters, and breakdowns."),
        ("5", "Investor report is generated", "The report script creates a clean Excel workbook with summary KPIs and definitions."),
    ]
    add_table(document, ["Step", "Stage", "Description"], workflow_rows, [0.55, 1.75, 4.2])

    document.add_page_break()
    document.add_heading("3. Main Project Files", level=1)
    file_rows = [
        ("src/events.js", "Defines event names, active user events, retention events, and allowed properties."),
        ("src/posthogAnalytics.js", "Central tracking and identify utility for PostHog."),
        ("src/locvmSchemaMapping.js", "Maps MongoDB user and locum documents into analytics-ready properties."),
        ("src/exportPosthogEvents.js", "Exports raw PostHog events into CSV."),
        ("src/generatePosthogMatrix.js", "Creates an internal metric matrix CSV from metric definitions."),
        ("src/generateInvestorReport.js", "Creates the investor-ready Excel workbook."),
        ("metrics/locvm_metrics_template.csv", "Excel-editable input file for defining metrics."),
        ("mock_data/mock_posthog_events.csv", "Sample event data for local testing."),
        ("POSTHOG_ANALYTICS.md", "Detailed technical documentation."),
    ]
    add_table(document, ["File", "Purpose"], file_rows, [2.35, 4.15])

    document.add_heading("4. Input Data Fields", level=1)
    document.add_paragraph("The project uses three broad input groups: PostHog event fields, LOCVM user fields, and LOCVM locum posting fields.")

    event_rows = [
        ("timestamp", "Event time used for reporting periods."),
        ("event", "PostHog event name, such as locum_applied or clinic_created_posting."),
        ("distinct_id / user_id", "User identifier used for unique user and retention calculations."),
        ("role", "User type such as physician or clinic."),
        ("locum_id / job_id / slug", "Locum posting identifiers."),
        ("specialty / province / city", "Breakdown fields for marketplace analysis."),
        ("pay", "Numeric pay value used for GMV and average pay calculations."),
        ("job_type", "Posting type such as FT, PT, or CONTRACT."),
        ("notification_type", "Notification category for notification click analysis."),
    ]
    add_table(document, ["PostHog Event Field", "How It Is Used"], event_rows, [2.2, 4.3])

    user_rows = [
        ("users._id", "Primary user identifier."),
        ("users.role", "Physician, Clinic, Admin, or User; normalized for analytics."),
        ("users.discoverySource", "Signup attribution and acquisition source."),
        ("users.medSpeciality", "Physician specialty."),
        ("users.medicalProvince / users.workAddress.province", "Province normalized for geographic reporting."),
        ("users.preferences.isLookingForLocums", "Supply activation and active physician context."),
        ("users.profile.profileCompletion", "Profile completion and onboarding reporting."),
        ("users.reservationsList.*", "Saved, applied, reserved, completed, and created locum activity."),
    ]
    add_table(document, ["LOCVM User Field", "How It Is Used"], user_rows, [2.8, 3.7])

    locum_rows = [
        ("locums._id", "Primary locum posting identifier."),
        ("locums.jobId / locums.slug", "Business-facing job identifiers."),
        ("locums.locumCreator", "Clinic or hospital user who created the posting."),
        ("locums.medSpeciality", "Required specialty for supply and demand analysis."),
        ("locums.jobType", "Posting type such as FT, PT, or CONTRACT."),
        ("locums.fullAddress.city / province", "Geographic breakdowns."),
        ("locums.locumPay", "Parsed into numeric pay for GMV and average pay."),
        ("locums.reservationId", "Filled/open status and marketplace liquidity."),
    ]
    add_table(document, ["LOCVM Locum Field", "How It Is Used"], locum_rows, [2.8, 3.7])

    document.add_page_break()
    document.add_heading("5. Outputs", level=1)
    output_rows = [
        ("Raw event export CSV", "exports/locvm_posthog_events_...csv", "Internal analyst file containing event-level data."),
        ("Metric matrix CSV", "exports/locvm_posthog_matrix_...csv", "Internal metric table with values, breakdowns, and date ranges."),
        ("Investor workbook", "reports/LOCVM_Investor_Report_...xlsx", "Business-facing Excel report with summary KPIs, KPI detail, raw event sample, and definitions."),
    ]
    add_table(document, ["Output", "Location", "Purpose"], output_rows, [1.5, 2.4, 2.6])

    document.add_heading("6. Investor Metrics Supported", level=1)
    metric_rows = [
        ("Marketplace Liquidity Score", "Filled locum postings / total locum postings."),
        ("Supply Activation Rate", "Active physicians / physician signups or total physicians, depending on data source."),
        ("Demand Activation Rate", "Clinics creating postings / clinic signups or total clinics."),
        ("Marketplace Balance Ratio", "Active physicians / open locum jobs."),
        ("Physician Engagement Depth", "Physician marketplace actions / active physicians."),
        ("Application Conversion Rate", "Locum applications / locum views."),
        ("Reservation Conversion Rate", "Reservations created / locum applications."),
        ("Locum Completion Rate", "Completed locums / reserved locums."),
        ("Marketplace GMV", "Sum of completed or reserved locum pay."),
        ("Active Users", "Unique users performing meaningful marketplace actions."),
    ]
    add_table(document, ["Metric", "Definition"], metric_rows, [2.4, 4.1])

    document.add_heading("7. Testing Commands", level=1)
    command_rows = [
        ("npm run check", "Checks JavaScript syntax for the project scripts."),
        ("npm run matrix:mock", "Generates the internal metric matrix using mock event data."),
        ("npm run report:mock", "Generates the investor-ready Excel workbook using mock event data."),
        ("npm run export -- --days=7", "Exports real PostHog events once credentials are configured."),
    ]
    add_table(document, ["Command", "Purpose"], command_rows, [2.5, 4.0])

    document.add_heading("8. Important Notes", level=1)
    add_bullets(document, [
        "This project does not modify MongoDB data.",
        "This project does not replace the existing GA4 dashboard.",
        "The actual LOCVM application still needs to call the tracking utility from production user flows.",
        "The investor workbook is the main business-facing output; CSV files are supporting analyst outputs.",
        "Mock data is included so the reporting flow can be tested before real PostHog credentials are configured.",
    ])

    add_footer(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build_document()
